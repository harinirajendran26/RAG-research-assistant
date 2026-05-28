import os
import re
import shutil
import time
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
from docx import Document as DocxDocument

load_dotenv()

CHROMA_BASE_PATH = "./chroma_db"
EMBEDDING_MODEL  = "BAAI/bge-base-en-v1.5"


# ── Versioned path helper ─────────────────────────────────────────────────────
def get_next_chroma_path() -> str:
    i = 1
    while True:
        path = f"{CHROMA_BASE_PATH}_v{i}"
        if not os.path.exists(path):
            return path
        i += 1

def get_latest_chroma_path() -> str:
    i = 1
    latest = None
    while True:
        path = f"{CHROMA_BASE_PATH}_v{i}"
        if os.path.exists(path):
            latest = path
            i += 1
        else:
            break
    return latest


# ── Model loaders ─────────────────────────────────────────────────────────────
def load_embeddings_model():
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

def load_reranker():
    return CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")

def load_llm():
    return ChatGroq(
        model="llama-3.3-70b-versatile",
        api_key=os.getenv("GROQ_API_KEY")
    )

def load_vectorstore(embeddings_model):
    latest = get_latest_chroma_path()
    if latest:
        return Chroma(
            persist_directory=latest,
            embedding_function=embeddings_model
        )
    return None


# ── DOCX loader ───────────────────────────────────────────────────────────────
def load_docx(file_path: str) -> list:
    """Load a .docx file and return as a list with one Document."""
    filename  = os.path.basename(file_path)
    doc       = DocxDocument(file_path)
    full_text = "\n\n".join([
        para.text for para in doc.paragraphs
        if len(para.text.strip()) > 20
    ])
    return [Document(
        page_content=full_text,
        metadata={"source": filename, "type": "docx", "section": "main"}
    )]


# ── TXT loader ────────────────────────────────────────────────────────────────
def load_txt(file_path: str) -> list:
    """Load a .txt file and return as a list with one Document."""
    filename = os.path.basename(file_path)
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"source": filename, "type": "txt", "section": "main"}
    )]


# ── Web page loader ───────────────────────────────────────────────────────────
def load_url(url: str) -> Document:
    """Scrape a web page and return as a Document."""
    try:
        headers  = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers, timeout=10)
        soup     = BeautifulSoup(response.text, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text  = soup.get_text(separator="\n")
        lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 30]
        clean = "\n".join(lines)

        title = soup.find("title")
        name  = title.get_text().strip()[:50] if title else url[:50]

        return Document(
            page_content=clean[:8000],
            metadata={"source": name, "url": url, "type": "web"}
        )
    except Exception as e:
        return Document(
            page_content=f"Failed to load {url}: {str(e)}",
            metadata={"source": url, "type": "web"}
        )


# ── PDF section extractor ─────────────────────────────────────────────────────
def extract_special_sections(full_text: str, filename: str) -> list:
    special          = []
    abstract_start   = full_text.find("Abstract\n")
    intro_start      = full_text.find("1 Introduction\n")
    background_start = full_text.find("2 Background\n")

    if abstract_start != -1 and intro_start != -1:
        abstract_text = full_text[abstract_start:intro_start].strip()
        lines = [l for l in abstract_text.split('\n')
                 if not l.strip().startswith(('∗','†','‡','31st','arXiv'))]
        abstract_text = '\n'.join(lines).strip()
        special.append(Document(
            page_content=abstract_text,
            metadata={"source": filename, "section": "abstract",
                      "page": 0, "type": "pdf"}
        ))

    if intro_start != -1 and background_start != -1:
        intro_text = full_text[intro_start:background_start].strip()
        special.append(Document(
            page_content=intro_text,
            metadata={"source": filename, "section": "introduction",
                      "page": 1, "type": "pdf"}
        ))
    return special


def clean_text(text: str) -> str:
    text = re.sub(r'\S+@\S+', '', text)
    text = re.sub(r'http\S+', '', text)
    lines = text.split('\n')
    return '\n'.join([
        l.strip() for l in lines
        if len(l.strip()) > 15 or
        (l.strip() and l.strip()[0].isdigit())
    ])


# ── Main ingestion pipeline ───────────────────────────────────────────────────
def process_sources(
    pdf_paths: list,
    urls: list,
    text_inputs: list,
    embeddings_model
) -> tuple:
    """
    Process PDFs, DOCX, TXT, URLs, and raw text inputs.
    Returns (summary dict, new vectorstore).
    """
    new_path   = get_next_chroma_path()
    all_chunks = []
    summary    = {}

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800, chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    # ── Process files (PDF, DOCX, TXT) ───────────────────────────────────────
    for file_path in pdf_paths:
        filename  = os.path.basename(file_path)
        extension = filename.lower().rsplit('.', 1)[-1]

        if extension == 'pdf':
            loader     = PyPDFLoader(file_path)
            pages      = loader.load()
            full_text  = "\n\n".join([p.page_content for p in pages])
            page_count = len(pages)

            special   = extract_special_sections(full_text, filename)

            bg_start  = full_text.find("2 Background\n")
            in_start  = full_text.find("1 Introduction\n")
            rem_start = bg_start if bg_start != -1 else \
                        in_start if in_start != -1 else 0

            remaining  = clean_text(full_text[rem_start:])
            base_docs  = [Document(
                page_content=remaining,
                metadata={"source": filename, "section": "main",
                          "page": 0, "type": "pdf"}
            )]

        elif extension == 'docx':
            base_docs  = load_docx(file_path)
            special    = []
            page_count = 1

        elif extension == 'txt':
            base_docs  = load_txt(file_path)
            special    = []
            page_count = 1

        else:
            print(f"Skipping unsupported file type: {filename}")
            continue

        # Chunk base documents
        regular_chunks = splitter.split_documents(base_docs)

        # Tag source on special chunks
        for chunk in special:
            chunk.metadata["source"] = filename

        file_chunks = special + regular_chunks
        all_chunks.extend(file_chunks)
        summary[filename] = {
            "type":   extension,
            "pages":  page_count,
            "chunks": len(file_chunks)
        }
        print(f"  {filename} ({extension}): {page_count} page(s) → {len(file_chunks)} chunks")

    # ── Process URLs ──────────────────────────────────────────────────────────
    for url in urls:
        doc    = load_url(url)
        chunks = splitter.split_documents([doc])
        for c in chunks:
            c.metadata.update({
                "source":  doc.metadata["source"],
                "url":     url,
                "type":    "web",
                "section": "main"
            })
        all_chunks.extend(chunks)
        summary[doc.metadata["source"]] = {
            "type": "url", "url": url,
            "chunks": len(chunks)
        }

    # ── Process raw text inputs ───────────────────────────────────────────────
    for i, text_input in enumerate(text_inputs):
        name   = f"note_{i+1}"
        chunks = splitter.split_documents([Document(
            page_content=text_input,
            metadata={"source": name, "type": "text", "section": "main"}
        )])
        all_chunks.extend(chunks)
        summary[name] = {"type": "text", "chunks": len(chunks)}

    # ── Embed everything ──────────────────────────────────────────────────────
    vectorstore = Chroma.from_documents(
        documents=all_chunks,
        embedding=embeddings_model,
        persist_directory=new_path
    )

    return summary, vectorstore


# ── Retrieval ─────────────────────────────────────────────────────────────────
def retrieve(
    question: str,
    vectorstore,
    reranker,
    k_fetch: int = 20,
    k_final: int = 4
) -> list:
    """
    Smart retrieval — routes by question type,
    re-ranks candidates with cross-encoder.
    """
    summary_keywords = ["main idea", "summary", "overview",
                        "about", "purpose", "what is this"]
    is_summary = any(kw in question.lower() for kw in summary_keywords)

    if is_summary:
        chunks  = vectorstore.similarity_search(
            question, k=1,
            filter={"section": {"$eq": "abstract"}}
        )
        chunks += vectorstore.similarity_search(
            question, k=1,
            filter={"section": {"$eq": "introduction"}}
        )
        if not chunks:
            chunks = vectorstore.similarity_search(question, k=k_fetch)
    else:
        chunks = vectorstore.similarity_search(question, k=k_fetch)

    if not chunks:
        return []

    # Re-rank with cross-encoder
    pairs  = [(question, doc.page_content) for doc in chunks]
    scores = reranker.predict(pairs)
    scored = sorted(zip(scores, chunks), key=lambda x: x[0], reverse=True)
    return [doc for _, doc in scored[:k_final]]


# ── Generation ────────────────────────────────────────────────────────────────
def generate(
    question: str,
    chunks: list,
    history: list,
    llm
) -> str:
    """Generate answer using retrieved chunks and conversation history."""
    context = "\n\n".join([
        f"[Source: {doc.metadata.get('source','?')} | "
        f"Type: {doc.metadata.get('type','?')} | "
        f"Section: {doc.metadata.get('section','?')}]\n"
        f"{doc.page_content}"
        for doc in chunks
    ])

    recent       = history[-4:] if len(history) > 4 else history
    history_text = "\n".join([
        f"{'User' if m['role']=='user' else 'Assistant'}: {m['content'][:200]}"
        for m in recent
    ])

    messages = [
        SystemMessage(content="""You are a helpful research assistant.
Answer using ONLY the provided context.
Always cite sources using (Source: name, type).
If the answer is not in context say 'I could not find this in the loaded sources.'
Be direct and specific — never give evasive answers."""),
        HumanMessage(content=
            f"Conversation history:\n{history_text}\n\n"
            f"Context:\n{context}\n\n"
            f"Question: {question}"
        )
    ]
    return llm.invoke(messages).content

# ── Streaming generation ──────────────────────────────────────────────────────
def generate_stream(
    question: str,
    chunks: list,
    history: list,
    llm
):
    """
    Same as generate() but yields tokens as they arrive.
    Returns a generator — use with st.write_stream() in Streamlit.
    """
    context = "\n\n".join([
        f"[Source: {doc.metadata.get('source','?')} | "
        f"Type: {doc.metadata.get('type','?')} | "
        f"Section: {doc.metadata.get('section','?')}]\n"
        f"{doc.page_content}"
        for doc in chunks
    ])

    recent       = history[-4:] if len(history) > 4 else history
    history_text = "\n".join([
        f"{'User' if m['role']=='user' else 'Assistant'}: {m['content'][:200]}"
        for m in recent
    ])

    messages = [
        SystemMessage(content="""You are a helpful research assistant.
Answer using ONLY the provided context.
Always cite sources using (Source: name, type).
If the answer is not in context say 'I could not find this in the loaded sources.'
Be direct and specific — never give evasive answers."""),
        HumanMessage(content=
            f"Conversation history:\n{history_text}\n\n"
            f"Context:\n{context}\n\n"
            f"Question: {question}"
        )
    ]

    # Stream tokens as they arrive
    full_response = ""
    for chunk in llm.stream(messages):
        token = chunk.content
        full_response += token
        yield token

# ── Cost estimator ────────────────────────────────────────────────────────────
def estimate_cost(text: str) -> dict:
    """
    Estimate token count and cost for Groq Llama 3.3 70B.
    Groq pricing: ~$0.59 per 1M input tokens, $0.79 per 1M output tokens.
    Rough estimate: 1 token ≈ 4 characters.
    """
    chars       = len(text)
    tokens      = chars // 4
    input_cost  = (tokens / 1_000_000) * 0.59
    return {
        "estimated_tokens": tokens,
        "estimated_cost_usd": round(input_cost, 6)
    }

# ── Full pipeline ─────────────────────────────────────────────────────────────
# ── Full pipeline with latency + cost tracking ────────────────────────────────
def answer(
    question: str,
    history: list,
    vectorstore,
    reranker,
    llm
) -> dict:
    if not vectorstore:
        return {
            "answer":   "No sources loaded. Please add PDFs, URLs, or text.",
            "sources":  [], "chunks": [],
            "metrics":  {}
        }

    # ── Track retrieval time ──────────────────────────────────────────────────
    t0     = time.time()
    chunks = retrieve(question, vectorstore, reranker)
    retrieval_time = round(time.time() - t0, 2)

    if not chunks:
        return {
            "answer":  "Could not find relevant information in loaded sources.",
            "sources": [], "chunks": [],
            "metrics": {"retrieval_time_s": retrieval_time}
        }

    # ── Build context for cost estimate ───────────────────────────────────────
    context      = "\n\n".join([doc.page_content for doc in chunks])
    cost_info    = estimate_cost(context + question)

    # ── Track generation time ─────────────────────────────────────────────────
    t1       = time.time()
    response = generate(question, chunks, history, llm)
    generation_time = round(time.time() - t1, 2)

    total_time = round(retrieval_time + generation_time, 2)

    # ── Build sources list ────────────────────────────────────────────────────
    sources = []
    seen    = set()
    for doc in chunks:
        src = doc.metadata.get('source', '?')
        typ = doc.metadata.get('type', '?')
        sec = doc.metadata.get('section', '?')
        key = f"{src}_{sec}"
        if key not in seen:
            sources.append({
                "source":  src,
                "type":    typ,
                "section": sec,
                "preview": doc.page_content[:120] + "..."
            })
            seen.add(key)

    # ── Build metrics dict ────────────────────────────────────────────────────
    metrics = {
        "retrieval_time_s":  retrieval_time,
        "generation_time_s": generation_time,
        "total_time_s":      total_time,
        "chunks_retrieved":  len(chunks),
        "estimated_tokens":  cost_info["estimated_tokens"],
        "estimated_cost_usd": cost_info["estimated_cost_usd"]
    }

    return {
        "answer":  response,
        "sources": sources,
        "chunks":  chunks,
        "metrics": metrics
    }

# ── Streaming full pipeline ───────────────────────────────────────────────────
def answer_stream(
    question: str,
    history: list,
    vectorstore,
    reranker,
    llm
) -> dict:
    """
    Same as answer() but returns a generator for streaming.
    Returns dict with stream generator + sources + metrics.
    """
    if not vectorstore:
        def empty(): yield "No sources loaded. Please add PDFs, URLs, or text."
        return {"stream": empty(), "sources": [], "chunks": [], "metrics": {}}

    # Retrieve with timing
    t0     = time.time()
    chunks = retrieve(question, vectorstore, reranker)
    retrieval_time = round(time.time() - t0, 2)

    if not chunks:
        def empty(): yield "Could not find relevant information in loaded sources."
        return {"stream": empty(), "sources": [], "chunks": [], "metrics": {}}

    context   = "\n\n".join([doc.page_content for doc in chunks])
    cost_info = estimate_cost(context + question)

    # Build sources
    sources = []
    seen    = set()
    for doc in chunks:
        src = doc.metadata.get('source', '?')
        typ = doc.metadata.get('type', '?')
        sec = doc.metadata.get('section', '?')
        key = f"{src}_{sec}"
        if key not in seen:
            sources.append({
                "source":  src,
                "type":    typ,
                "section": sec,
                "preview": doc.page_content[:120] + "..."
            })
            seen.add(key)

    # Return stream generator + metadata
    return {
        "stream":  generate_stream(question, chunks, history, llm),
        "sources": sources,
        "chunks":  chunks,
        "metrics": {
            "retrieval_time_s":   retrieval_time,
            "chunks_retrieved":   len(chunks),
            "estimated_tokens":   cost_info["estimated_tokens"],
            "estimated_cost_usd": cost_info["estimated_cost_usd"]
        }
    }

# ── Utility ───────────────────────────────────────────────────────────────────
def get_loaded_sources(vectorstore) -> list:
    if not vectorstore:
        return []
    results = vectorstore._collection.get()
    sources = {}
    for meta in results['metadatas']:
        if meta and 'source' in meta:
            src = meta['source']
            typ = meta.get('type', '?')
            if src not in sources:
                sources[src] = typ
    return [{"name": k, "type": v} for k, v in sorted(sources.items())]